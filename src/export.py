"""Export utilities for research reports."""

import io
import json
import logging
import re
import zipfile
from datetime import datetime
from typing import Optional

import markdown2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

logger = logging.getLogger(__name__)


def markdown_to_html(markdown_content: str) -> str:
    """Convert markdown to HTML with styling."""
    html_content = markdown2.markdown(
        markdown_content,
        extras=[
            "fenced-code-blocks",
            "tables",
            "header-ids",
            "strike",
            "task_list",
        ],
    )

    # Wrap in styled HTML document
    styled_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 1.6;
            color: #1a1a2e;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
        }}

        h1 {{
            color: #007A75;
            border-bottom: 3px solid #007A75;
            padding-bottom: 10px;
            font-size: 28px;
        }}

        h2 {{
            color: #1a1a2e;
            border-bottom: 1px solid #e0e0e0;
            padding-bottom: 8px;
            margin-top: 30px;
            font-size: 22px;
        }}

        h3 {{
            color: #333;
            margin-top: 25px;
            font-size: 18px;
        }}

        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}

        th {{
            background-color: #007A75;
            color: white;
        }}

        tr:nth-child(even) {{
            background-color: #f9f9f9;
        }}

        blockquote {{
            border-left: 4px solid #007A75;
            margin: 20px 0;
            padding: 10px 20px;
            background-color: #f8fafa;
            font-style: italic;
        }}

        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Fira Code', monospace;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}

        a {{
            color: #007A75;
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        ul, ol {{
            margin: 15px 0;
            padding-left: 30px;
        }}

        li {{
            margin: 8px 0;
        }}

        .source-citation {{
            font-size: 0.9em;
            color: #666;
        }}

        @media print {{
            body {{
                max-width: none;
                padding: 20px;
            }}

            h1, h2, h3 {{
                page-break-after: avoid;
            }}

            table, blockquote {{
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
    return styled_html


def export_to_pdf(markdown_content: str, title: str = "Research Report") -> bytes:
    """Export markdown content to PDF bytes."""
    logger.info("Generating PDF: title=%s, content_length=%d", title, len(markdown_content))
    try:
        from weasyprint import HTML

        html_content = markdown_to_html(markdown_content)
        pdf_bytes = HTML(string=html_content).write_pdf()
        logger.info("PDF generated successfully: %d bytes", len(pdf_bytes))
        return pdf_bytes
    except ImportError:
        logger.error("weasyprint not installed")
        raise ImportError(
            "weasyprint is required for PDF export. "
            "Install it with: pip install weasyprint"
        )
    except Exception as e:
        logger.error("PDF generation failed: %s", e, exc_info=True)
        raise RuntimeError(f"PDF generation failed: {e}")


def _add_hyperlink(paragraph, text, url):
    """Add a hyperlink run to a paragraph.

    Creates a proper OOXML hyperlink element so the link is clickable in Word.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    run_text = OxmlElement("w:t")
    run_text.text = text
    # Preserve leading/trailing spaces
    run_text.set(qn("xml:space"), "preserve")
    new_run.append(run_text)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


# Pattern for inline markdown: bold, italic, and links.
# Order matters — check bold before italic since ** contains *.
_INLINE_RE = re.compile(
    r"(\*\*(.+?)\*\*)"         # bold **text**
    r"|(\*(.+?)\*)"            # italic *text*
    r"|(\[([^\]]+)\]\(([^)]+)\))"  # link [text](url)
)


def _add_formatted_runs(paragraph, text):
    """Parse inline markdown (bold, italic, links) and add runs to *paragraph*.

    Plain text segments become normal runs.  ``**bold**`` becomes a bold run,
    ``*italic*`` becomes an italic run, and ``[text](url)`` becomes a
    clickable hyperlink.
    """
    last_end = 0
    for m in _INLINE_RE.finditer(text):
        # Add any plain text before this match
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        if m.group(2) is not None:
            # Bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(4) is not None:
            # Italic
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(6) is not None:
            # Link
            link_text = m.group(6)
            link_url = m.group(7)
            _add_hyperlink(paragraph, link_text, link_url)

        last_end = m.end()

    # Trailing plain text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _is_table_separator(line: str) -> bool:
    """Return True if *line* is a markdown table separator row (e.g. |---|---|)."""
    stripped = line.strip().strip("|").strip()
    return bool(stripped) and all(
        c in "-: " for c in stripped.replace("|", "")
    )


def _parse_table_row(line: str) -> list[str]:
    """Split a markdown table row into cell strings, stripping outer pipes."""
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def export_to_docx(markdown_content: str, title: str = "Research Report") -> bytes:
    """Export markdown content to DOCX bytes."""
    logger.info("Generating DOCX: title=%s, content_length=%d", title, len(markdown_content))
    doc = Document()

    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add generation date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # Parse markdown and add content
    lines = markdown_content.split("\n")
    in_code_block = False
    code_content = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_content)
                code_para = doc.add_paragraph(code_text)
                code_para.style = "No Spacing"
                for run in code_para.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Handle headings
        if line.startswith("# "):
            doc.add_heading(line[2:], 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], 3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], 4)
        # Handle blockquotes
        elif line.startswith("> "):
            quote_para = doc.add_paragraph()
            quote_para.style = "Quote"
            _add_formatted_runs(quote_para, line[2:])

        # ----- Markdown table detection -----
        # A table starts with a pipe-delimited row followed by a separator row.
        elif "|" in line and (i + 1 < len(lines)) and _is_table_separator(lines[i + 1]):
            header_cells = _parse_table_row(line)
            i += 1  # skip separator row

            # Collect data rows
            data_rows: list[list[str]] = []
            while i + 1 < len(lines) and "|" in lines[i + 1] and not _is_table_separator(lines[i + 1]):
                i += 1
                data_rows.append(_parse_table_row(lines[i]))

            num_cols = len(header_cells)
            table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
            table.style = "Table Grid"

            # Header row
            for col_idx, cell_text in enumerate(header_cells):
                if col_idx < num_cols:
                    cell = table.rows[0].cells[col_idx]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    run = para.add_run(cell_text)
                    run.bold = True

            # Data rows
            for row_idx, row_cells in enumerate(data_rows, start=1):
                for col_idx, cell_text in enumerate(row_cells):
                    if col_idx < num_cols and row_idx < len(table.rows):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = ""
                        _add_formatted_runs(cell.paragraphs[0], cell_text)

            logger.debug("Added table: %d cols, %d data rows", num_cols, len(data_rows))

        # Handle nested bullet points (indented with 2+ spaces)
        elif re.match(r"^( {2,}|\t+)[-*] ", line):
            indent_match = re.match(r"^(\s+)[-*] (.*)", line)
            if indent_match:
                text = indent_match.group(2)
                indent_level = len(indent_match.group(1))
                # 2-3 spaces or 1 tab = level 2, 4+ spaces or 2+ tabs = level 3
                if indent_level >= 4:
                    style = "List Bullet 3"
                else:
                    style = "List Bullet 2"
                para = doc.add_paragraph(style=style)
                _add_formatted_runs(para, text)

        # Handle bullet points (top-level)
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(para, line[2:])
        # Handle numbered lists
        elif line and line[0].isdigit() and ". " in line[:4]:
            text = line.split(". ", 1)[1] if ". " in line else line
            para = doc.add_paragraph(style="List Number")
            _add_formatted_runs(para, text)
        # Handle horizontal rules
        elif line.startswith("---"):
            doc.add_paragraph("─" * 50)
        # Handle regular paragraphs
        elif line.strip():
            para = doc.add_paragraph()
            _add_formatted_runs(para, line)

        i += 1

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.getvalue()
    logger.info("DOCX generated successfully: %d bytes", len(docx_bytes))
    return docx_bytes


def create_asset_bundle(
    report_markdown: str,
    raw_findings: Optional[dict] = None,
    title: str = "research_report",
) -> bytes:
    """Create a ZIP bundle with all research assets."""
    logger.info("Creating asset bundle: title=%s", title)
    buffer = io.BytesIO()

    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        # Add the main report as markdown
        zf.writestr(f"{title}.md", report_markdown)

        # Add HTML version
        html_content = markdown_to_html(report_markdown)
        zf.writestr(f"{title}.html", html_content)

        # Try to add PDF version
        try:
            pdf_bytes = export_to_pdf(report_markdown, title)
            zf.writestr(f"{title}.pdf", pdf_bytes)
        except Exception as e:
            logger.warning("PDF generation failed, skipping: %s", e)

        # Add DOCX version
        try:
            docx_bytes = export_to_docx(report_markdown, title)
            zf.writestr(f"{title}.docx", docx_bytes)
        except Exception as e:
            logger.warning("DOCX generation failed, skipping: %s", e)

        # Add raw findings if provided
        if raw_findings:
            # Create raw_data directory in zip
            for agent_name, findings in raw_findings.items():
                if isinstance(findings, dict):
                    content = json.dumps(findings, indent=2, default=str)
                    zf.writestr(f"raw_data/{agent_name}.json", content)
                else:
                    zf.writestr(f"raw_data/{agent_name}.txt", str(findings))

        # Add metadata
        metadata = {
            "generated_at": datetime.now().isoformat(),
            "title": title,
            "files_included": [
                f"{title}.md",
                f"{title}.html",
                f"{title}.pdf",
                f"{title}.docx",
            ],
        }
        if raw_findings:
            metadata["raw_data_files"] = list(raw_findings.keys())

        zf.writestr("metadata.json", json.dumps(metadata, indent=2))

    buffer.seek(0)
    logger.info("Asset bundle created: %d bytes", len(buffer.getvalue()))
    return buffer.getvalue()


def get_filename_from_title(title: str, extension: str = "") -> str:
    """Generate a safe filename from a title."""
    # Remove special characters and replace spaces
    safe_title = "".join(
        c if c.isalnum() or c in " -_" else "" for c in title
    ).strip()
    safe_title = safe_title.replace(" ", "-").lower()

    # Add date prefix
    date_prefix = datetime.now().strftime("%Y-%m-%d")
    filename = f"{date_prefix}-{safe_title}"

    if extension:
        filename = f"{filename}.{extension.lstrip('.')}"

    return filename
