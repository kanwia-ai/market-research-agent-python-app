"""Tests for export functionality."""

import json
import zipfile
from io import BytesIO

from docx import Document

from src.export import (
    _is_table_separator,
    _parse_table_row,
    create_asset_bundle,
    export_to_docx,
    get_filename_from_title,
    markdown_to_html,
)


class TestMarkdownToHtml:
    """Tests for markdown to HTML conversion."""

    def test_converts_basic_markdown(self):
        """Test basic markdown conversion."""
        md = "# Hello World\n\nThis is a paragraph."
        html = markdown_to_html(md)
        assert "<h1" in html
        assert "Hello World" in html
        assert "<p>" in html

    def test_includes_styling(self):
        """Test that HTML includes CSS styling."""
        md = "# Test"
        html = markdown_to_html(md)
        assert "<style>" in html
        assert "font-family" in html

    def test_handles_tables(self):
        """Test table conversion."""
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        html = markdown_to_html(md)
        assert "<table>" in html

    def test_handles_code_blocks(self):
        """Test code block conversion."""
        md = "```python\nprint('hello')\n```"
        html = markdown_to_html(md)
        assert "<code" in html or "<pre" in html


class TestExportToDocx:
    """Tests for DOCX export."""

    def test_returns_bytes(self):
        """Test that export returns bytes."""
        md = "# Test Report\n\nSome content."
        result = export_to_docx(md, "Test")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_creates_valid_docx(self):
        """Test that output is valid DOCX (ZIP format)."""
        md = "# Test\n\n- Item 1\n- Item 2"
        result = export_to_docx(md)
        # DOCX files are ZIP archives
        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            # DOCX must contain these files
            assert "word/document.xml" in zf.namelist()

    def test_handles_headings(self):
        """Test heading conversion."""
        md = "# H1\n## H2\n### H3"
        result = export_to_docx(md)
        assert len(result) > 0

    def test_handles_lists(self):
        """Test list conversion."""
        md = "- Bullet 1\n- Bullet 2\n\n1. Numbered 1\n2. Numbered 2"
        result = export_to_docx(md)
        assert len(result) > 0

    def test_handles_blockquotes(self):
        """Test blockquote conversion."""
        md = '> This is a quote\n> — Author'
        result = export_to_docx(md)
        assert len(result) > 0


class TestCreateAssetBundle:
    """Tests for asset bundle creation."""

    def test_returns_zip_bytes(self):
        """Test that bundle returns ZIP bytes."""
        md = "# Report"
        result = create_asset_bundle(md)
        assert isinstance(result, bytes)

        # Verify it's a valid ZIP
        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            assert len(zf.namelist()) > 0

    def test_includes_markdown_file(self):
        """Test that bundle includes markdown."""
        md = "# My Report"
        result = create_asset_bundle(md, title="test")

        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            assert "test.md" in zf.namelist()
            content = zf.read("test.md").decode("utf-8")
            assert "My Report" in content

    def test_includes_html_file(self):
        """Test that bundle includes HTML."""
        md = "# Report"
        result = create_asset_bundle(md, title="test")

        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            assert "test.html" in zf.namelist()

    def test_includes_docx_file(self):
        """Test that bundle includes DOCX."""
        md = "# Report"
        result = create_asset_bundle(md, title="test")

        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            assert "test.docx" in zf.namelist()

    def test_includes_metadata(self):
        """Test that bundle includes metadata."""
        md = "# Report"
        result = create_asset_bundle(md, title="test")

        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            assert "metadata.json" in zf.namelist()
            metadata = json.loads(zf.read("metadata.json"))
            assert "generated_at" in metadata
            assert metadata["title"] == "test"

    def test_includes_raw_findings(self):
        """Test that raw findings are included."""
        md = "# Report"
        findings = {
            "CommunityMapper": {"communities": ["reddit", "twitter"]},
            "VoiceMiner": {"quotes": ["quote1", "quote2"]},
        }
        result = create_asset_bundle(md, raw_findings=findings, title="test")

        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            names = zf.namelist()
            assert "raw_data/CommunityMapper.json" in names
            assert "raw_data/VoiceMiner.json" in names


class TestGetFilenameFromTitle:
    """Tests for filename generation."""

    def test_generates_safe_filename(self):
        """Test safe filename generation."""
        result = get_filename_from_title("My Report", "pdf")
        assert ".pdf" in result
        assert "/" not in result
        assert "\\" not in result

    def test_includes_date_prefix(self):
        """Test that date is included."""
        result = get_filename_from_title("Test")
        # Should have YYYY-MM-DD format at start
        assert result[4] == "-"
        assert result[7] == "-"

    def test_handles_special_characters(self):
        """Test handling of special characters."""
        result = get_filename_from_title("Test: Report! @#$%", "md")
        assert ":" not in result
        assert "!" not in result
        assert "@" not in result

    def test_lowercase_output(self):
        """Test that output is lowercase."""
        result = get_filename_from_title("UPPERCASE Test")
        # After the date prefix
        assert "uppercase" in result or "test" in result


def _load_docx(docx_bytes: bytes) -> Document:
    """Helper: load DOCX bytes back into a Document for inspection."""
    return Document(BytesIO(docx_bytes))


class TestInlineFormatting:
    """Tests for inline bold, italic, and link formatting in DOCX export."""

    def test_bold_text_in_paragraph(self):
        """Bold markdown (**text**) should produce a bold run."""
        md = "This has **bold words** in it."
        doc = _load_docx(export_to_docx(md))
        # Find the paragraph with our text (skip title/date/spacer)
        body_paras = [p for p in doc.paragraphs if "bold words" in p.text]
        assert len(body_paras) == 1
        para = body_paras[0]
        bold_runs = [r for r in para.runs if r.bold]
        assert any(r.text == "bold words" for r in bold_runs)

    def test_italic_text_in_paragraph(self):
        """Italic markdown (*text*) should produce an italic run."""
        md = "This has *italic words* in it."
        doc = _load_docx(export_to_docx(md))
        body_paras = [p for p in doc.paragraphs if "italic words" in p.text]
        assert len(body_paras) == 1
        para = body_paras[0]
        italic_runs = [r for r in para.runs if r.italic]
        assert any(r.text == "italic words" for r in italic_runs)

    def test_link_in_paragraph(self):
        """Markdown links [text](url) should create a hyperlink in the document XML."""
        md = "Visit [Example](https://example.com) for more."
        result = export_to_docx(md)
        # Hyperlinks are stored in the document XML relationships, not as normal runs.
        # Verify the relationship exists by checking the raw DOCX XML.
        buf = BytesIO(result)
        with zipfile.ZipFile(buf, "r") as zf:
            rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
            assert "https://example.com" in rels_xml
            doc_xml = zf.read("word/document.xml").decode("utf-8")
            assert "Example" in doc_xml

    def test_mixed_inline_formatting(self):
        """Paragraph with bold, italic, and plain text together."""
        md = "Start **bold** then *italic* and plain."
        doc = _load_docx(export_to_docx(md))
        body_paras = [p for p in doc.paragraphs if "bold" in p.text and "italic" in p.text]
        assert len(body_paras) == 1
        para = body_paras[0]
        bold_runs = [r for r in para.runs if r.bold]
        italic_runs = [r for r in para.runs if r.italic]
        assert any(r.text == "bold" for r in bold_runs)
        assert any(r.text == "italic" for r in italic_runs)

    def test_inline_formatting_in_list_items(self):
        """Bold/italic should also work inside bullet list items."""
        md = "- Item with **bold** text\n- Item with *italic* text"
        doc = _load_docx(export_to_docx(md))
        # List items are paragraphs too
        bold_paras = [p for p in doc.paragraphs if "bold" in p.text]
        assert len(bold_paras) >= 1
        bold_runs = [r for r in bold_paras[0].runs if r.bold]
        assert any(r.text == "bold" for r in bold_runs)

    def test_plain_text_no_formatting(self):
        """A paragraph with no markdown formatting should have a single plain run."""
        md = "Just plain text here."
        doc = _load_docx(export_to_docx(md))
        body_paras = [p for p in doc.paragraphs if "Just plain text here" in p.text]
        assert len(body_paras) == 1
        para = body_paras[0]
        assert len(para.runs) == 1
        assert not para.runs[0].bold
        assert not para.runs[0].italic


class TestTableParsing:
    """Tests for markdown table -> DOCX table conversion."""

    def test_simple_table(self):
        """A basic 2-col table should produce a DOCX table."""
        md = "| Name | Value |\n|------|-------|\n| A    | 1     |\n| B    | 2     |"
        doc = _load_docx(export_to_docx(md))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3  # header + 2 data rows
        assert len(table.columns) == 2

    def test_table_header_is_bold(self):
        """Table header cells should have bold text."""
        md = "| Col1 | Col2 |\n|------|------|\n| x    | y    |"
        doc = _load_docx(export_to_docx(md))
        header_row = doc.tables[0].rows[0]
        for cell in header_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        assert run.bold, f"Header cell '{run.text}' should be bold"

    def test_table_data_content(self):
        """Table data cells should contain the correct text."""
        md = "| Animal | Sound |\n|--------|-------|\n| Dog    | Woof  |\n| Cat    | Meow  |"
        doc = _load_docx(export_to_docx(md))
        table = doc.tables[0]
        assert table.rows[1].cells[0].text == "Dog"
        assert table.rows[1].cells[1].text == "Woof"
        assert table.rows[2].cells[0].text == "Cat"
        assert table.rows[2].cells[1].text == "Meow"

    def test_table_with_inline_formatting(self):
        """Inline bold in table cells should render as bold runs."""
        md = "| Feature | Status |\n|---------|--------|\n| Auth    | **Done** |"
        doc = _load_docx(export_to_docx(md))
        data_cell = doc.tables[0].rows[1].cells[1]
        runs = data_cell.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        assert any("Done" in r.text for r in bold_runs)

    def test_three_column_table(self):
        """A 3-column table should have 3 columns."""
        md = "| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |"
        doc = _load_docx(export_to_docx(md))
        assert len(doc.tables) == 1
        assert len(doc.tables[0].columns) == 3

    def test_table_separator_detection(self):
        """Helper _is_table_separator should identify separator rows."""
        assert _is_table_separator("|---|---|")
        assert _is_table_separator("| --- | --- |")
        assert _is_table_separator("|:---|---:|")
        assert not _is_table_separator("| data | more |")
        assert not _is_table_separator("not a table")

    def test_parse_table_row(self):
        """Helper _parse_table_row should split pipe-delimited cells."""
        assert _parse_table_row("| A | B | C |") == ["A", "B", "C"]
        assert _parse_table_row("|A|B|") == ["A", "B"]

    def test_table_does_not_break_surrounding_content(self):
        """Content before and after a table should still render."""
        md = "Before table.\n\n| H1 | H2 |\n|----|----|\n| a  | b  |\n\nAfter table."
        doc = _load_docx(export_to_docx(md))
        texts = [p.text for p in doc.paragraphs]
        assert any("Before table" in t for t in texts)
        assert any("After table" in t for t in texts)
        assert len(doc.tables) == 1


class TestNestedLists:
    """Tests for nested (indented) bullet list support."""

    def test_two_space_indent_uses_list_bullet_2(self):
        """A line starting with '  - ' should use 'List Bullet 2' style."""
        md = "- Top item\n  - Nested item"
        doc = _load_docx(export_to_docx(md))
        nested_paras = [p for p in doc.paragraphs if "Nested item" in p.text]
        assert len(nested_paras) == 1
        assert nested_paras[0].style.name == "List Bullet 2"

    def test_four_space_indent_uses_list_bullet_3(self):
        """A line starting with '    - ' should use 'List Bullet 3' style."""
        md = "- Top\n    - Deep nested"
        doc = _load_docx(export_to_docx(md))
        nested_paras = [p for p in doc.paragraphs if "Deep nested" in p.text]
        assert len(nested_paras) == 1
        assert nested_paras[0].style.name == "List Bullet 3"

    def test_top_level_still_uses_list_bullet(self):
        """A non-indented bullet should still use 'List Bullet'."""
        md = "- Top level"
        doc = _load_docx(export_to_docx(md))
        bullet_paras = [p for p in doc.paragraphs if "Top level" in p.text]
        assert len(bullet_paras) == 1
        assert bullet_paras[0].style.name == "List Bullet"

    def test_nested_with_asterisk(self):
        """Nested items using * instead of - should also work."""
        md = "* Parent\n  * Child"
        doc = _load_docx(export_to_docx(md))
        child_paras = [p for p in doc.paragraphs if "Child" in p.text]
        assert len(child_paras) == 1
        assert child_paras[0].style.name == "List Bullet 2"

    def test_nested_items_have_inline_formatting(self):
        """Nested list items should also support bold/italic."""
        md = "- Top\n  - **Bold nested** item"
        doc = _load_docx(export_to_docx(md))
        nested_paras = [p for p in doc.paragraphs if "Bold nested" in p.text]
        assert len(nested_paras) == 1
        bold_runs = [r for r in nested_paras[0].runs if r.bold]
        assert any("Bold nested" in r.text for r in bold_runs)

    def test_mixed_nesting_levels(self):
        """Multiple nesting levels in sequence."""
        md = "- Level 1\n  - Level 2\n    - Level 3\n- Back to 1"
        doc = _load_docx(export_to_docx(md))
        l1 = [p for p in doc.paragraphs if p.text == "Level 1"]
        l2 = [p for p in doc.paragraphs if p.text == "Level 2"]
        l3 = [p for p in doc.paragraphs if p.text == "Level 3"]
        back = [p for p in doc.paragraphs if p.text == "Back to 1"]
        assert l1[0].style.name == "List Bullet"
        assert l2[0].style.name == "List Bullet 2"
        assert l3[0].style.name == "List Bullet 3"
        assert back[0].style.name == "List Bullet"
